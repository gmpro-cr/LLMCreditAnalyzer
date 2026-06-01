"""
Credit Risk Summary Generator.

Produces a decision-grade 8-section credit risk summary.
Python pre-computes ALL numbers, ratios, flags, and trends.
The LLM performs ONLY analytical interpretation and risk judgment.

Sections:
  1. Borrower Overview
  2. Financial Analysis
  3. Liquidity & Cash Flow Assessment
  4. Key Risk Drivers
  5. Mitigating Factors
  6. Red Flags
  7. Overall Risk Assessment (LOW / MODERATE / HIGH)
  8. Confidence Level (HIGH / MEDIUM / LOW)
"""
import os
import re
import logging
import httpx
from typing import Dict, Any, List, Optional
from credit_risk_engine import build_context

logger = logging.getLogger(__name__)


# ── LLM Providers ─────────────────────────────────────────────────────────────

def _claude_call(prompt: str) -> str:
    """Call Claude via the `claude` CLI (uses Claude Code Pro subscription)."""
    import subprocess
    system = (
        "You are a senior credit risk analyst at a commercial bank. "
        "Write formal, precise credit analysis in Markdown. "
        "Use ## for section headers. Use **bold** for key figures and risk ratings. "
        "Reference ONLY the exact numbers provided to you. "
        "NEVER fabricate data. Output ONLY the final credit risk summary in clean Markdown."
    )
    try:
        result = subprocess.run(
            ["claude", "-p", f"{system}\n\n{prompt}"],
            capture_output=True, text=True, timeout=300,
        )
        if result.returncode != 0:
            logger.error(f"Claude CLI error: {result.stderr}")
            return ""
        return result.stdout.strip()
    except Exception as e:
        logger.error(f"Claude CLI call error: {e}")
        return ""


def _call_llm(prompt: str) -> str:
    """Call configured LLM. Tries memo provider first, then fallback."""
    provider = os.getenv("MEMO_PROVIDER", "gemini").lower()
    if provider == "claude":
        result = _claude_call(prompt)
        if result:
            return result
        logger.warning("Claude failed, falling back to Gemini")
        result = _gemini_call(prompt)
        if result:
            return result
        logger.warning("Gemini failed, falling back to Ollama")
    elif provider == "gemini":
        result = _gemini_call(prompt)
        if result:
            return result
        logger.warning("Gemini failed, falling back to Ollama")
    return _ollama_call(prompt)


def _gemini_call(prompt: str) -> str:
    try:
        from google import genai
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            return ""
        client = genai.Client(api_key=api_key)
        response = client.models.generate_content(
            model="models/gemini-flash-lite-latest",
            contents=prompt,
        )
        return response.text.strip()
    except Exception as e:
        logger.error(f"Gemini call error: {e}")
        return ""


def _ollama_call(prompt: str) -> str:
    ollama_url = os.getenv("OLLAMA_URL", "http://localhost:11434")
    ollama_model = os.getenv("OLLAMA_MODEL", "qwen3:8b")
    try:
        resp = httpx.post(
            f"{ollama_url}/v1/chat/completions",
            json={
                "model": ollama_model,
                "messages": [
                    {
                        "role": "system",
                        "content": (
                            "You are a senior credit risk analyst at a commercial bank. "
                            "Write formal, precise credit analysis in Markdown. "
                            "Use ## for section headers. Use **bold** for key figures and risk ratings. "
                            "Reference ONLY the exact numbers provided to you. "
                            "NEVER fabricate data. NEVER use vague language like 'the company is doing well'. "
                            "Instead use precise statements like 'EBITDA margin declined from X% to Y%'. "
                            "Prioritize cash flow over profit signals. "
                            "Do NOT include any thinking, reasoning traces, or <think> tags in your output. "
                            "Output ONLY the final credit risk summary in clean Markdown."
                        ),
                    },
                    {"role": "user", "content": prompt},
                ],
                "temperature": 0.3,
                "stream": False,
            },
            timeout=300.0,
        )
        resp.raise_for_status()
        result = resp.json()["choices"][0]["message"]["content"].strip()
        # Strip any thinking traces that qwen3 might emit
        result = re.sub(r'<think>.*?</think>', '', result, flags=re.DOTALL).strip()
        return result
    except Exception as e:
        logger.error(f"Ollama call error: {e}")
        return ""


# ── 8-Section Credit Risk Summary Prompt ──────────────────────────────────────

def _build_credit_risk_prompt(context_text: str, company_name: str) -> str:
    """
    Build the LLM prompt for the 8-section credit risk summary.
    Optimized for smaller models (qwen3:8b): very explicit structure,
    pre-computed data injected, and tight format constraints.
    """
    return f"""You are a senior credit risk analyst at a commercial bank.
Generate a detailed, decision-grade credit risk summary using ONLY the pre-computed data below.

STRICT RULES:
- Use ONLY the numbers provided. Do NOT assume or fabricate any values.
- Where data says "null — cannot be computed", write "Data insufficient to assess."
- Prioritize cash flow signals over reported profits.
- Every risk claim MUST cite the specific metric value as evidence.
- Be concise but insight-heavy. No storytelling.
- Format output as clean Markdown with ## headers for each section.

{context_text}

═══════════════════════════════════════════════════════════════════════
Write EXACTLY these 8 sections. Follow the format precisely:
═══════════════════════════════════════════════════════════════════════

## 1. Borrower Overview

Write 3-4 sentences: Who is {company_name}? What does the company do? Industry, scale, key business segments. Use ABOUT and EXTERNAL INTELLIGENCE sections above.

## 2. Financial Analysis

Write 3-4 paragraphs covering key trends (NOT a raw data dump):
- Revenue trajectory: growth/decline pattern, CAGR, drivers
- Profitability: EBITDA margin trend across years, PAT movement, why improving or deteriorating
- Leverage: debt trajectory, Debt/Equity trend, interest burden changes
- Balance sheet strength: net worth accretion, asset quality

Reference specific year-over-year numbers from the FINANCIAL DATA BY YEAR section.

## 3. Liquidity & Cash Flow Assessment

Write 2-3 paragraphs:
- Cash flow from operations: adequate vs inadequate for debt servicing?
- Compare CFO with PAT — are profits backed by actual cash generation?
- DSCR and interest coverage adequacy
- Free cash flow generation capability
- Working capital pressures (if inferable from the data)

## 4. Key Risk Drivers

List 3-5 major risks. For EACH risk:
- State the risk clearly
- Provide the specific metric or data point as evidence
- Use this format:

**Risk 1: [Risk Title]**
Evidence: [specific metric value from pre-computed data]

Use the AUTO-DETECTED RED FLAGS section as your primary input, but add analytical context.

## 5. Mitigating Factors

List the counterbalancing strengths from the AUTO-DETECTED MITIGATING FACTORS section.
For each, add one sentence of analytical context explaining why it matters for credit risk.
Avoid generic statements.

## 6. Red Flags

List ALL red flags from the AUTO-DETECTED RED FLAGS section as a bullet list.
For each, include the severity tag and evidence exactly as provided.
If no red flags were detected, write: "No critical red flags identified by automated screening."

## 7. Overall Risk Assessment

State ONE of: **LOW RISK** / **MODERATE RISK** / **HIGH RISK**

Then write 4-5 lines of justification referencing:
- The trend direction of key ratios
- Cash flow adequacy
- Leverage sustainability
- Any red flags or mitigants that shift the rating

## 8. Confidence Level

State ONE of: **HIGH** / **MEDIUM** / **LOW**

Then explain based on:
- Data completeness (how many years? all metrics available?)
- Data recency (is it current or stale?)
- Availability of qualitative inputs (research brief, credit ratings, management info)

Use the DATA QUALITY section above for this assessment.

═══════════════════════════════════════════════════════════════════════
CRITICAL: Output ONLY the 8 sections above. No preamble, no conclusion,
no "here is the summary" text. Start directly with ## 1. Borrower Overview
═══════════════════════════════════════════════════════════════════════"""


# ── Main Entry Points ─────────────────────────────────────────────────────────

def generate_credit_risk_summary(
    financials: Dict,
    ratios: Dict,
    company_name: str,
    covenants: list = None,
    research_brief: str = "",
) -> str:
    """
    Generate full 8-section credit risk summary.
    Python pre-computes all data → LLM performs analytical interpretation.
    """
    # Step 1: Build pre-computed context
    ctx = build_context(financials, ratios, research_brief, company_name)

    # Hard stop if data is insufficient
    if ctx.get("abort"):
        return f"# Credit Risk Assessment — BLOCKED\n\n**{ctx['abort_reason']}**\n\nData quality issues:\n" + \
               "\n".join(f"- {w}" for w in ctx["validation"].get("warnings", []))

    # Step 2: Build LLM prompt with pre-computed context
    prompt = _build_credit_risk_prompt(ctx["context_text"], company_name)

    # Step 3: Call LLM for analytical interpretation
    logger.info(f"[CreditRisk] Generating 8-section summary for {company_name}")
    summary = _call_llm(prompt)

    if not summary:
        # Fallback: return structured context without LLM narrative
        return _fallback_summary(ctx, company_name)

    # Step 4: Prepend header
    header = f"# Credit Risk Summary — {company_name}\n\n"
    header += f"*Generated by CreditGuard AI | Confidence: {ctx['validation']['confidence_level']}*\n\n---\n\n"

    return header + summary


def generate_cam_memo(
    financials: Dict,
    ratios: Dict,
    company_name: str,
    covenants: list = None,
    research_brief: str = "",
) -> str:
    """
    Backward-compatible entry point — redirects to generate_credit_risk_summary,
    then injects the deterministic financial spread + ratio-covenant tables.
    """
    from cam_tables import inject_into_memo
    memo = generate_credit_risk_summary(
        financials, ratios, company_name, covenants, research_brief
    )
    return inject_into_memo(memo, financials, ratios)


def _fallback_summary(ctx: Dict, company_name: str) -> str:
    """Generate a minimal summary from pre-computed data when LLM is unavailable."""
    lines = [
        f"# Credit Risk Summary — {company_name}",
        f"\n*LLM unavailable — showing pre-computed data only*\n",
        "---",
        "\n## Pre-Computed Financial Metrics\n",
    ]

    for yd in ctx.get("yearly_data", []):
        lines.append(f"### {yd['year']}")
        if yd.get("revenue"): lines.append(f"- Revenue: ₹{yd['revenue']:,.0f} Cr")
        if yd.get("ebitda"): lines.append(f"- EBITDA: ₹{yd['ebitda']:,.0f} Cr")
        if yd.get("pat"): lines.append(f"- PAT: ₹{yd['pat']:,.0f} Cr")
        if yd.get("total_debt"): lines.append(f"- Total Debt: ₹{yd['total_debt']:,.0f} Cr")
        lines.append("")

    if ctx.get("red_flags"):
        lines.append("## Red Flags\n")
        for rf in ctx["red_flags"]:
            lines.append(f"- **[{rf['severity']}]** {rf['flag']}")
            lines.append(f"  - {rf['evidence']}")
        lines.append("")

    if ctx.get("mitigants"):
        lines.append("## Mitigating Factors\n")
        for m in ctx["mitigants"]:
            lines.append(f"- {m}")

    lines.append(f"\n## Confidence Level: {ctx['validation']['confidence_level']}")

    return "\n".join(lines)


# ── DOCX Export ───────────────────────────────────────────────────────────────

def export_to_docx(memo_content: str, company_name: str, output_path: str) -> None:
    """Export memo Markdown to a Word .docx file."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        for para in doc.paragraphs:
            para.clear()
        for line in memo_content.split("\n"):
            if line.startswith("# "):
                p = doc.add_heading(line[2:], level=0)
            elif line.startswith("## "):
                p = doc.add_heading(line[3:], level=1)
            elif line.startswith("### "):
                p = doc.add_heading(line[4:], level=2)
            elif line.startswith("| "):
                doc.add_paragraph(line)
            elif line.strip():
                p = doc.add_paragraph()
                parts = re.split(r"\*\*(.+?)\*\*", line)
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 == 1:
                        run.bold = True
            else:
                doc.add_paragraph("")
        doc.save(output_path)
    except ImportError:
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(memo_content)
